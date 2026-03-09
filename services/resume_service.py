import hashlib
import io
import os
import re
from typing import List, Optional
import json
import logging


from dotenv import load_dotenv
from fastapi import HTTPException
from openai import OpenAI
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from datetime import date 

from models.llm_run import LlmRun
from models.resume import Resume
from models.resume_classification import ResumeClassification
from models.resume_keyword import ResumeKeyword
from schemas.resume_llm import (
    ResumeClassificationResult,
    ResumeKeywordItem,
    ResumeKeywordResult,
)
from services.prompt.resume.classify_prompt_v2 import (
    PROMPT_VERSION_CLASSIFY,
    CLASSIFY_SYSTEM_PROMPT,
    build_classify_user_prompt,
)
from services.prompt.resume.keyword_prompt_v2 import (
    PROMPT_VERSION_KEYWORD,
    KEYWORD_SYSTEM_PROMPT,
    build_keyword_user_prompt,
)

from models.resume_structured import ResumeStructured
from schemas.resume_structured import ResumeStructuredResult
from services.prompt.resume.structure_prompt_v1 import (
    PROMPT_VERSION_STRUCTURE,
    STRUCTURE_SYSTEM_PROMPT,
    build_structure_user_prompt,
)

logger = logging.getLogger(__name__)

load_dotenv()

DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

def update_resume_status(
    db: Session,
    resume: Resume,
    status_value: str,
    error_message: Optional[str] = None,
) -> Resume:
    resume.resume_status = status_value
    resume.resume_error_message = (error_message or "")[:255] or None
    db.add(resume)
    db.commit()
    db.refresh(resume)

    logger.info(
        "RESUME_STATUS_UPDATED resume_id=%s status=%s error=%s",
        resume.resume_id,
        resume.resume_status,
        resume.resume_error_message,
    )
    return resume

class ResumeFileError(Exception):
    def __init__(self, detail: str, status_code: int = 400):
        super().__init__(detail)
        self.detail = detail
        self.status_code = status_code


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

SECTION_HINTS = [
    "학력", "교육", "경력", "경험", "프로젝트", "기술", "스킬", "역량",
    "자격증", "수상", "성과", "대외활동", "포트폴리오", "자기소개", "요약",
]

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(01[016789])[-\s]?\d{3,4}[-\s]?\d{4}")
DATE_RE = re.compile(r"(19|20)\d{2}[./-]\d{1,2}")

# career_summary 값 정규화/검증
def normalize_career_summary(value: Optional[str]) -> Optional[str]:
    if not value:
        return None

    text = re.sub(r"\s+", " ", value).strip()
    if not text:
        return None

    # 1년 미만/인턴/신입 표현은 모두 신입으로 통일
    newbie_keywords = [
        "신입",
        "인턴",
        "연수",
        "교육",
        "부트캠프",
        "취업준비",
        "취준",
    ]
    if any(keyword in text for keyword in newbie_keywords):
        return "신입"

    # "총 5년", "경력 5년", "5년 3개월", "약 5년" -> "5년"
    year_match = re.search(r"(\d+)\s*년", text)
    if year_match:
        years = int(year_match.group(1))
        if years <= 0:
            return "신입"
        return f"{years}년"

    # 개월만 있으면 1년 미만으로 보고 신입 처리
    month_match = re.search(r"(\d+)\s*개월", text)
    if month_match:
        months = int(month_match.group(1))
        if months < 12:
            return "신입"

    return None

def _parse_year_month(value: Optional[str]) -> Optional[tuple[int, int]]:
    if not value:
        return None

    text = value.strip()

    m = re.search(r"((?:19|20)\d{2})[./-](\d{1,2})", text)
    if m:
        year = int(m.group(1))
        month = int(m.group(2))
        if 1 <= month <= 12:
            return year, month

    m = re.search(r"((?:19|20)\d{2})", text)
    if m:
        year = int(m.group(1))
        return year, 1

    return None


def _month_diff(start: tuple[int, int], end: tuple[int, int]) -> int:
    sy, sm = start
    ey, em = end
    return max(0, (ey - sy) * 12 + (em - sm) + 1)


def calculate_career_summary_from_experiences(experiences: List) -> str:
    total_months = 0

    for exp in experiences or []:
        exp_type = getattr(exp, "experience_type", None)
        count_as_career = getattr(exp, "count_as_career", False)

        # FULL_TIME / CONTRACT 이면서 count_as_career=true 인 경우만 포함
        if not count_as_career:
            continue

        if exp_type not in {"FULL_TIME", "CONTRACT"}:
            continue

        start = _parse_year_month(getattr(exp, "start_date", None))
        end_raw = getattr(exp, "end_date", None)

        if not start:
            continue

        if end_raw and any(word in end_raw for word in ["재직", "현재", "근무중"]):
            today = date.today()
            end = (today.year, today.month)
        else:
            end = _parse_year_month(end_raw)

        if not end:
            continue

        total_months += _month_diff(start, end)

    years = total_months // 12

    if years <= 0:
        return "신입"

    return f"{years}년"


def is_probable_resume(text: str):
    t = (text or "").strip()
    reasons = []
    score = 0

    if len(t) < 500:
        reasons.append("텍스트 길이가 너무 짧음(<500자)")
    else:
        score += 1

    hits = [k for k in SECTION_HINTS if k in t]
    if len(hits) >= 2:
        score += 2
    elif len(hits) == 1:
        score += 1
        reasons.append("이력서 섹션 키워드가 1개만 탐지됨")
    else:
        reasons.append("이력서 섹션 키워드가 거의 없음")

    if EMAIL_RE.search(t):
        score += 1
    else:
        reasons.append("이메일 형식이 탐지되지 않음")

    if PHONE_RE.search(t):
        score += 1

    if DATE_RE.search(t):
        score += 1

    ok = score >= 4
    return ok, reasons, score



def detect_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return "PDF"
    if ext == ".docx":
        return "DOCX"

    raise HTTPException(
        status_code=400,
        detail="PDF/DOCX만 업로드 가능합니다.",
    )


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def extract_pdf_text(data: bytes) -> str:
    try:
        import fitz
    except Exception as e:
        raise RuntimeError("pymupdf가 필요합니다.") from e

    doc = None
    try:
        doc = fitz.open(stream=data, filetype="pdf")

        # 실제 비밀번호 입력이 필요한 PDF만 차단
        if getattr(doc, "needs_pass", False):
            raise ResumeFileError("암호가 설정된 PDF는 업로드할 수 없습니다.", 400)

        # 암호화 플래그만 있고 읽기 가능한 경우가 있어 추가 확인
        if getattr(doc, "is_encrypted", False):
            auth_result = doc.authenticate("")
            if auth_result == 0 and getattr(doc, "needs_pass", False):
                raise ResumeFileError("암호가 설정된 PDF는 업로드할 수 없습니다.", 400)

        parts: List[str] = []
        for page in doc:
            page_text = page.get_text("text") or ""
            if page_text.strip():
                parts.append(page_text)

        text = "\n".join(parts).strip()

        if not text:
            raise ResumeFileError(
                "PDF에서 텍스트를 추출하지 못했습니다. 스캔본, 이미지형 PDF, 또는 손상된 PDF일 수 있습니다.",
                400,
            )

        return text

    except ResumeFileError:
        raise
    except Exception as e:
        raise ResumeFileError(f"PDF 파싱 실패: {e}", 400) from e
    finally:
        if doc is not None:
            doc.close()


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx가 필요합니다.") from e

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ResumeFileError(
            "DOCX를 열 수 없습니다. 암호로 보호되었거나 손상된 파일일 수 있습니다.",
            400,
        ) from e

    try:
        parts: List[str] = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()

        if not text:
            raise ResumeFileError(
                "DOCX에서 텍스트를 추출하지 못했습니다. 비어 있거나 읽을 수 없는 문서입니다.",
                400,
            )

        return text

    except ResumeFileError:
        raise
    except Exception as e:
        raise ResumeFileError(f"DOCX 파싱 실패: {e}", 400) from e


def extract_text_from_upload(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return extract_pdf_text(data)

    if ext == ".docx":
        return extract_docx_text(data)

    raise HTTPException(status_code=400, detail="지원하지 않는 파일 형식입니다.")


def get_client() -> OpenAI:
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY가 없습니다.")
    return OpenAI(api_key=api_key)


def save_llm_run_success(
    db: Session,
    stage: str,
    model: str,
    prompt_version: str,
) -> LlmRun:
    row = LlmRun(
        llm_stage=stage,
        llm_model=model,
        llm_prompt_version=prompt_version,
        llm_status="SUCCESS",
        error_code=None,
        error_message=None,
    )
    db.add(row)
    db.flush()
    return row


def save_llm_run_failed(
    db: Session,
    stage: str,
    model: str,
    prompt_version: str,
    error_code: str,
    error_message: str,
) -> None:
    row = LlmRun(
        llm_stage=stage,
        llm_model=model,
        llm_prompt_version=prompt_version,
        llm_status="FAILED",
        error_code=error_code,
        error_message=(error_message or "")[:255] or None,
    )
    db.add(row)
    db.flush()


def create_resume_record(
    db: Session,
    user_id: int,
    original_filename: str,
    data: bytes,
) -> Resume:
    file_type = detect_file_type(original_filename)

    try:
        extracted_text = normalize_text(
            extract_text_from_upload(original_filename, data)
        )
    except ResumeFileError as e:
        raise HTTPException(status_code=e.status_code, detail=e.detail) from e

    if not extracted_text:
        raise HTTPException(status_code=400, detail="텍스트를 추출하지 못했습니다.")

    ok, reasons, score = is_probable_resume(extracted_text)
    if not ok:
        reason_msg = ", ".join(reasons[:2])
        raise HTTPException(
            status_code=400,
            detail=f"이 문서는 이력서로 보이지 않습니다. (score={score}) 사유: {reason_msg}",
        )

    if len(extracted_text) > 80000:
        extracted_text = extracted_text[:80000] + "\n\n[TRUNCATED]"

    file_hash = sha256_bytes(data)

    existing = (
        db.query(Resume)
        .filter(
            Resume.user_id == user_id,
            Resume.resume_sha256 == file_hash,
        )
        .first()
    )

    if existing:
        raise HTTPException(
            status_code=409,
            detail="이미 업로드한 이력서입니다.",
        )

    resume = Resume(
        user_id=user_id,
        resume_file_name=original_filename,
        resume_file_type=file_type,
        resume_file_path=None,
        resume_file_size=len(data),
        resume_extracted_text=extracted_text,
        resume_sha256=sha256_bytes(data),
        resume_status="UPLOADED",
        resume_error_message=None,
    )
    try:
        db.add(resume)
        db.commit()
        db.refresh(resume)
        return resume

    except IntegrityError as e:
        db.rollback()
        raise HTTPException(
            status_code=409,
            detail="이미 업로드한 이력서입니다.",
        ) from e

    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"이력서 저장 실패: {e}",
        ) from e


def get_resume_by_id(db: Session, resume_id: int) -> Resume:
    resume = db.query(Resume).filter(Resume.resume_id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="이력서를 찾을 수 없습니다.")
    return resume


def get_resume_analysis_result(db: Session, resume_id: int):
    resume = get_resume_by_id(db, resume_id)

    classification = (
        db.query(ResumeClassification)
        .filter(ResumeClassification.resume_id == resume_id)
        .first()
    )

    keywords = (
        db.query(ResumeKeyword)
        .filter(ResumeKeyword.resume_id == resume_id)
        .order_by(ResumeKeyword.keyword_id.asc())
        .all()
    )

    structured = (
        db.query(ResumeStructured)
        .filter(ResumeStructured.resume_id == resume_id)
        .first()
    )

    return {
        "resume": resume,
        "classification": classification,
        "keywords": keywords,
        "structured": structured,
    }


def classify_resume_llm(
    resume_text: str,
    model: str = DEFAULT_MODEL,
) -> ResumeClassificationResult:
    client = get_client()

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": CLASSIFY_SYSTEM_PROMPT},
            {"role": "user", "content": build_classify_user_prompt(resume_text)},
        ],
        text_format=ResumeClassificationResult,
        truncation="auto",
    )

    if resp.output_parsed is None:
        raise RuntimeError("이력서 분류 파싱 실패")

    return resp.output_parsed

def analyze_resume_structure_llm(
    resume_text: str,
    job_family: Optional[str],
    job_role: Optional[str],
    model: str = DEFAULT_MODEL,
) -> ResumeStructuredResult:
    client = get_client()

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": STRUCTURE_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": build_structure_user_prompt(
                    resume_text=resume_text,
                    job_family=job_family,
                    job_role=job_role,
                ),
            },
        ],
        text_format=ResumeStructuredResult,
        truncation="auto",
    )

    if resp.output_parsed is None:
        raise RuntimeError("이력서 구조화 분석 파싱 실패")

    return resp.output_parsed


def analyze_resume_keywords_llm(
    structured_payload: dict,
    job_family: Optional[str],
    job_role: Optional[str],
    model: str = DEFAULT_MODEL,
) -> ResumeKeywordResult:
    client = get_client()

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": KEYWORD_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": build_keyword_user_prompt(
                    structured_json=json.dumps(structured_payload, ensure_ascii=False, indent=2),
                    job_family=job_family,
                    job_role=job_role,
                ),
            },
        ],
        text_format=ResumeKeywordResult,
        truncation="auto",
    )

    if resp.output_parsed is None:
        raise RuntimeError("이력서 키워드 분석 파싱 실패")

    return resp.output_parsed


def dedupe_keywords(items: List[ResumeKeywordItem]) -> List[ResumeKeywordItem]:
    result: List[ResumeKeywordItem] = []
    seen = set()

    for item in items:
        keyword = (item.keyword or "").strip()
        if not keyword:
            continue

        key = (keyword.lower(), item.keyword_type)
        if key in seen:
            continue

        seen.add(key)
        result.append(item)

    return result


def analyze_saved_resume(
    db: Session,
    resume_id: int,
    model: str = DEFAULT_MODEL,
) -> None:
    resume = get_resume_by_id(db, resume_id)

    classification_row = (
        db.query(ResumeClassification)
        .filter(ResumeClassification.resume_id == resume_id)
        .first()
    )

    structure_row = (
        db.query(ResumeStructured)
        .filter(ResumeStructured.resume_id == resume_id)
        .first()
    )

    keyword_exists = (
        db.query(ResumeKeyword)
        .filter(ResumeKeyword.resume_id == resume_id)
        .first()
    )

    # 이미 분석 완료된 경우 상태 보정
    if classification_row and structure_row and keyword_exists:
        if resume.resume_status != "DONE":
            update_resume_status(db, resume, "KEYWORDS_DONE")
        return

    if not classification_row:
        try:
            update_resume_status(db, resume, "CLASSIFYING")

            classification_result = classify_resume_llm(
                resume_text=resume.resume_extracted_text,
                model=model,
            )

            if not classification_result.is_resume:
                raise HTTPException(status_code=400, detail="이 문서는 이력서가 아닙니다.")

            classify_run = save_llm_run_success(
                db=db,
                stage="RESUME_CLASSIFY",
                model=model,
                prompt_version=PROMPT_VERSION_CLASSIFY,
            )

            classification_row = ResumeClassification(
                resume_id=resume.resume_id,
                llm_id=classify_run.llm_id,
                class_job_family=classification_result.job_family,
                class_job_role=classification_result.job_role,
                class_evidence=[x.model_dump() for x in classification_result.evidence],
            )
            db.add(classification_row)
            db.commit()
            db.refresh(classification_row)

        except Exception as e:
            db.rollback()
            save_llm_run_failed(
                db=db,
                stage="RESUME_CLASSIFY",
                model=model,
                prompt_version=PROMPT_VERSION_CLASSIFY,
                error_code=type(e).__name__,
                error_message=str(e),
            )
            db.commit()

            update_resume_status(db, resume, "FAILED", str(e))
            raise HTTPException(status_code=500, detail=f"이력서 분류 실패: {e}") from e

    if not structure_row:
        try:
            update_resume_status(db, resume, "STRUCTURING")

            structure_result = analyze_resume_structure_llm(
                resume_text=resume.resume_extracted_text,
                job_family=classification_row.class_job_family if classification_row else None,
                job_role=classification_row.class_job_role if classification_row else None,
                model=model,
            )

            structure_run = save_llm_run_success(
                db=db,
                stage="RESUME_STRUCTURE",
                model=model,
                prompt_version=PROMPT_VERSION_STRUCTURE,
            )

            career_summary = calculate_career_summary_from_experiences(
                structure_result.experiences
            )

            if not career_summary:
                career_summary = normalize_career_summary(structure_result.career_summary)

            structure_row = ResumeStructured(
                resume_id=resume.resume_id,
                llm_id=structure_run.llm_id,
                structured_position=structure_result.position,
                structured_career_summary=career_summary,
                structured_skills=structure_result.skills,
                structured_educations=[x.model_dump() for x in structure_result.educations],
                structured_experiences=[x.model_dump() for x in structure_result.experiences],
                structured_projects=[x.model_dump() for x in structure_result.projects],
                structured_certificates=[x.model_dump() for x in structure_result.certificates],
            )
            db.add(structure_row)
            db.commit()
            db.refresh(structure_row)

        except Exception as e:
            db.rollback()
            save_llm_run_failed(
                db=db,
                stage="RESUME_STRUCTURE",
                model=model,
                prompt_version=PROMPT_VERSION_STRUCTURE,
                error_code=type(e).__name__,
                error_message=str(e),
            )
            db.commit()

            update_resume_status(db, resume, "FAILED", str(e))
            raise HTTPException(status_code=500, detail=f"이력서 구조화 분석 실패: {e}") from e

    if not keyword_exists:
        try:
            if structure_row is None:
                structure_row = (
                    db.query(ResumeStructured)
                    .filter(ResumeStructured.resume_id == resume_id)
                    .first()
                )

            if structure_row is None:
                raise RuntimeError("키워드 분석 전에 구조화 결과를 찾을 수 없습니다.")

            update_resume_status(db, resume, "KEYWORDS_EXTRACTING")

            structured_payload = build_structured_payload(structure_row)

            keyword_result = analyze_resume_keywords_llm(
                structured_payload=structured_payload,
                job_family=classification_row.class_job_family if classification_row else None,
                job_role=classification_row.class_job_role if classification_row else None,
                model=model,
            )

            deduped = dedupe_keywords(keyword_result.keywords)

            keyword_run = save_llm_run_success(
                db=db,
                stage="RESUME_KEYWORD",
                model=model,
                prompt_version=PROMPT_VERSION_KEYWORD,
            )

            for item in deduped:
                db.add(
                    ResumeKeyword(
                        resume_id=resume.resume_id,
                        llm_id=keyword_run.llm_id,
                        keyword_keyword=item.keyword,
                        keyword_type=item.keyword_type,
                        keyword_evidence=[x.model_dump() for x in item.evidence],
                    )
                )

            db.commit()

            update_resume_status(db, resume, "KEYWORDS_DONE")

        except Exception as e:
            db.rollback()
            save_llm_run_failed(
                db=db,
                stage="RESUME_KEYWORD",
                model=model,
                prompt_version=PROMPT_VERSION_KEYWORD,
                error_code=type(e).__name__,
                error_message=str(e),
            )
            db.commit()

            update_resume_status(db, resume, "FAILED", str(e))
            raise HTTPException(status_code=500, detail=f"이력서 키워드 분석 실패: {e}") from e
        
def delete_resume(db: Session, resume_id: int) -> None:
    resume = db.query(Resume).filter(Resume.resume_id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="이력서를 찾을 수 없습니다.")

    try:
        db.delete(resume)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"이력서 삭제 실패: {e}") from e
    
def build_structured_payload(structured_row: ResumeStructured) -> dict:
    return {
        "position": structured_row.structured_position,
        "career_summary": structured_row.structured_career_summary,
        "skills": structured_row.structured_skills or [],
        "educations": structured_row.structured_educations or [],
        "experiences": structured_row.structured_experiences or [],
        "projects": structured_row.structured_projects or [],
        "certificates": structured_row.structured_certificates or [],
    }
